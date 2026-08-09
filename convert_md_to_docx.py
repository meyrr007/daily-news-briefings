import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import re
import os

def md_to_docx(md_filepath: str, output_docx_path: str):
    """
    Converts a Markdown file into a beautifully styled Word (.docx) document.
    """
    doc = docx.Document()
    
    # Page Margins (0.75 in)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
    with open(md_filepath, "r", encoding="utf-8") as f:
        lines = f.readlines()
        
    for line in lines:
        raw_line = line.rstrip("\n")
        stripped = raw_line.strip()
        
        if not stripped:
            continue
            
        # Horizontal Rule (---)
        if stripped in ["---", "***", "___"]:
            p = doc.add_paragraph()
            p_format = p.paragraph_format
            p_format.space_before = Pt(4)
            p_format.space_after = Pt(8)
            pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/></w:pBdr>')
            p._p.get_or_add_pPr().append(pBdr)
            continue
            
        # Headings
        if stripped.startswith("# "):
            h = doc.add_heading(level=1)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(4)
            run = h.add_run(stripped[2:])
            run.font.name = 'Calibri'
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
            continue
        elif stripped.startswith("## "):
            h = doc.add_heading(level=2)
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(3)
            run = h.add_run(stripped[3:])
            run.font.name = 'Calibri'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x00, 0x40, 0x80)
            continue
        elif stripped.startswith("### "):
            h = doc.add_heading(level=3)
            h.paragraph_format.space_before = Pt(8)
            h.paragraph_format.space_after = Pt(2)
            run = h.add_run(stripped[4:])
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
            continue
        elif stripped.startswith("#### "):
            h = doc.add_heading(level=4)
            h.paragraph_format.space_before = Pt(6)
            h.paragraph_format.space_after = Pt(2)
            run = h.add_run(stripped[5:])
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            continue

        # Bullet points
        is_bullet = False
        if stripped.startswith("* ") or stripped.startswith("- "):
            is_bullet = True
            content = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
        else:
            content = stripped
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(4)

        # Parse Bold (**text**) and Italic (*text*) formatting in line
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', content)
        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = p.add_run(part[1:-1])
                run.italic = True
            else:
                run = p.add_run(part)
                
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    doc.save(output_docx_path)
    print(f"[+] Converted '{md_filepath}' -> '{output_docx_path}'")
    return output_docx_path

if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        md_file = sys.argv[1]
        out_file = md_file.rsplit('.', 1)[0] + '.docx'
        md_to_docx(md_file, out_file)
    else:
        md_to_docx(r"D:\papers\Michael_Bailey_Resume_Yahoo_Sr_TPM.md", r"D:\papers\Michael_Bailey_Resume_Yahoo_Sr_TPM.docx")
